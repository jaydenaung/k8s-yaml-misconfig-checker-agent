"""
build_docs.py — One-shot generator for the KubeSentinel technical documentation.

Run:  python build_docs.py
Output: KubeSentinel_Technical_Documentation.docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def shade_cell(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x10, 0x2A, 0x43) if level == 0 else RGBColor(0x1F, 0x3A, 0x68)
    return h


def add_para(doc, text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    return p


def add_code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    rPr = run._element.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), "Courier New")
    rFonts.set(qn("w:hAnsi"), "Courier New")
    rPr.append(rFonts)
    # Light grey shading via paragraph properties
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F2F2F2")
    pPr.append(shd)
    return p


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shade_cell(hdr[i], "1F3A68")
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)
    return table


def main():
    doc = Document()

    # Default style
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # ── Title page ─────────────────────────────────────────────────────────────
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("KubeSentinel")
    run.bold = True
    run.font.size = Pt(36)
    run.font.color.rgb = RGBColor(0x10, 0x2A, 0x43)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("AI-Powered Kubernetes Security Platform")
    run.font.size = Pt(16)
    run.italic = True
    run.font.color.rgb = RGBColor(0x4A, 0x4A, 0x4A)

    doc.add_paragraph()
    sub2 = doc.add_paragraph()
    sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub2.add_run("Technical Documentation")
    run.bold = True
    run.font.size = Pt(20)

    tag = doc.add_paragraph()
    tag.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = tag.add_run("Detect.  Reason.  Fix.")
    run.italic = True
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()
    doc.add_paragraph()

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for line in [
        "Author: Jayden Aung",
        "License: Apache License 2.0",
        "Default Model: claude-sonnet-4-6",
        "Document Version: 1.0",
    ]:
        run = meta.add_run(line + "\n")
        run.font.size = Pt(11)

    doc.add_page_break()

    # ── Table of contents (manual list) ────────────────────────────────────────
    add_heading(doc, "Table of Contents", level=1)
    toc_items = [
        "1.  Executive Summary",
        "2.  Purpose and Value Proposition",
        "3.  What KubeSentinel Actually Does",
        "4.  High-Level Architecture",
        "5.  Component Deep Dive",
        "       5.1  CLI Entry Point (agent.py)",
        "       5.2  Static Analyzer (analyzer.py)",
        "       5.3  Agentic Layer (claude_agent.py)",
        "       5.4  Tool Layer (tools.py)",
        "       5.5  Reporter (reporter.py)",
        "       5.6  Suppressor (suppressor.py)",
        "       5.7  Web Server (server.py and web/)",
        "       5.8  Database Schema (web/database.py)",
        "       5.9  Background Scanner (web/scanner.py)",
        "       5.10 Scheduler (web/scheduler.py)",
        "6.  The Two-Phase Scan -> Patch Design",
        "7.  Compound Risk Correlation Engine",
        "8.  Token-Efficient Cluster Fingerprinting",
        "9.  Static Check Catalogue (K8S-001 through K8S-014)",
        "10. Demonstration Guide",
        "11. Starting and Stopping the Server",
        "12. CLI Reference",
        "13. Configuration and Environment",
        "14. Security, Trust, and Limitations",
        "15. Project Layout",
        "16. Roadmap",
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.left_indent = Inches(0.15)
    doc.add_page_break()

    # ── 1. Executive Summary ───────────────────────────────────────────────────
    add_heading(doc, "1. Executive Summary", level=1)
    add_para(doc,
        "KubeSentinel is an on-premises, AI-powered Kubernetes security platform that goes "
        "beyond traditional static scanners. While conventional tools follow the loop "
        "\"Ingest -> Detect -> Surface -> Human decides -> Human acts\", KubeSentinel closes "
        "the loop with \"Observe -> Reason -> Patch -> Explain -> Human approves\". The agent "
        "does not merely flag that a setting is wrong; it correlates that finding with CVE "
        "data, RBAC exposure, and network policy gaps to produce a compound risk score, then "
        "generates corrected YAML patches on demand."
    )
    add_para(doc,
        "The platform is built around the Anthropic Claude API (default model: "
        "claude-sonnet-4-6) using the tool_use interface. A senior-Kubernetes-security-architect "
        "system prompt drives an iterative reasoning loop in which Claude decides which tools "
        "to call, in what order, based on what it has already observed. The static-check "
        "engine (14 deterministic checks) runs without any AI dependency, while the AI layer "
        "adds logic-level reasoning, supply-chain context, RBAC probing, and patch generation."
    )
    add_para(doc,
        "KubeSentinel ships three user-facing surfaces: a CLI (agent.py), a multi-user FastAPI "
        "web dashboard (server.py), and a GitHub Actions PR-scanner workflow. Findings persist "
        "into a local SQLite database; all data stays inside the user's environment with no SaaS "
        "dependency."
    )

    # ── 2. Purpose and Value Proposition ───────────────────────────────────────
    add_heading(doc, "2. Purpose and Value Proposition", level=1)
    add_para(doc, "Why this tool exists", bold=True)
    add_para(doc,
        "Kubernetes is operationally hostile to security: a single misconfigured pod can "
        "compromise an entire cluster, yet most teams ship YAML faster than they can review it. "
        "Existing scanners (kube-bench, kubesec, kubeaudit, Trivy, Polaris) excel at flagging "
        "isolated issues but suffer two structural weaknesses:"
    )
    add_bullets(doc, [
        "They report findings in isolation. A privileged pod, a wildcard ServiceAccount, and a "
        "namespace without NetworkPolicy each look survivable. Together they form an exploit chain.",
        "They stop at detection. Engineers still own the diagnosis, the YAML correction, and the "
        "explanation to their reviewers."
    ])
    add_para(doc, "What KubeSentinel adds", bold=True)
    add_bullets(doc, [
        "Cross-signal reasoning: correlates CVE + misconfiguration + RBAC + network signals per "
        "pod into proven exploit chains (CMP-001 through CMP-004).",
        "Runtime verification: kubectl auth can-i --as impersonation confirms what a "
        "ServiceAccount can actually do, not just what its RBAC theoretically allows.",
        "On-demand patch generation: a separate post-scan loop calls suggest_patch to produce "
        "minimal, corrected YAML snippets with one-sentence explanations.",
        "Air-gap friendly: full static analysis (14 checks, ~50 ms on a typical manifest) "
        "requires no API key, no network access, and no telemetry.",
        "CI/CD native: exits with code 2 on CRITICAL findings; ships a ready-to-copy GitHub "
        "Actions workflow that posts findings as PR comments and can block merges."
    ])

    # ── 3. What KubeSentinel actually does ─────────────────────────────────────
    add_heading(doc, "3. What KubeSentinel Actually Does", level=1)
    add_para(doc, "Inputs accepted:", bold=True)
    add_bullets(doc, [
        "A single Kubernetes YAML manifest file (multi-document files supported).",
        "A directory of manifests (recursively scanned, excluding .git, venv, node_modules, etc.).",
        "A Helm chart directory (Chart.yaml detected; rendered via 'helm template' before analysis).",
        "A live cluster, reached through an uploaded kubeconfig (web dashboard) or the current "
        "shell's kubectl context (CLI)."
    ])
    add_para(doc, "Outputs produced:", bold=True)
    add_bullets(doc, [
        "A list of findings with severity, check ID, context, detail, remediation, and (for AI "
        "findings) attack_scenario and telco_relevance fields.",
        "Compound risk findings (CMP-001..CMP-004) when 2+ signals correlate on the same pod.",
        "ServiceAccount probe results (SAP-001) confirming dangerous runtime API access.",
        "Optional AI-generated YAML patches for every finding, addressable by --patch.",
        "A Markdown report, a JSON dump, a GitHub PR comment, or persisted SQLite rows visible "
        "through the web dashboard.",
        "An exit code: 0 = clean, 1 = error, 2 = CRITICAL detected (suitable for CI gating)."
    ])
    add_para(doc, "What it never does", bold=True)
    add_bullets(doc, [
        "Modify your cluster. KubeSentinel uses only read-only kubectl verbs and 'auth can-i' "
        "impersonation checks. It never patches, deletes, scales, or rolls anything.",
        "Send your YAML to a public endpoint without your knowledge. The Anthropic API call is "
        "the only outbound dependency, and only when AI mode is enabled.",
        "Replace human review. AI-generated findings and patches are marked [AI] and must be "
        "independently verified before being applied."
    ])

    # ── 4. Architecture ────────────────────────────────────────────────────────
    add_heading(doc, "4. High-Level Architecture", level=1)
    add_para(doc,
        "KubeSentinel has three user-facing entry points (CLI, Web Dashboard, GitHub Actions) "
        "feeding into a shared agent layer that talks to Claude over the Anthropic tool_use "
        "API. Claude drives the analysis iteratively; it calls scan tools (load_manifest, "
        "render_helm_chart, run_check, query_cluster, lookup_image_cves, probe_service_account, "
        "scan_cluster_images, report_finding, finish) in whatever order makes sense given what "
        "the previous tool returned. Scan results land either in stdout (CLI) or in the SQLite "
        "database (web) along with file uploads under data/."
    )
    add_para(doc, "Layered view (top to bottom):", bold=True)
    add_code(doc,
        "User Interfaces\n"
        "  - CLI (python agent.py)\n"
        "  - Web Dashboard (FastAPI on port 8000)\n"
        "  - GitHub Actions PR Scanner (.github/workflows/kubesentinel.yml)\n"
        "\n"
        "FastAPI Server (server.py)\n"
        "  - REST API + Jinja2 dashboard\n"
        "  - Session auth (bcrypt, secret cookie in data/.secret_key)\n"
        "  - APScheduler (SQLAlchemy jobstore) for recurring cluster scans\n"
        "  - BackgroundTask runners for async scan execution\n"
        "\n"
        "AI Agent Layer (claude_agent.py)\n"
        "  - Scan loop:  analyze_with_agent / analyze_cluster_with_agent\n"
        "  - Patch loop: generate_patches_for_findings\n"
        "  - MAX_ITERATIONS = 25 per loop\n"
        "  - Two tool sets: scan tools (no suggest_patch), patch tools (suggest_patch + finish)\n"
        "\n"
        "Tool Layer (tools.py)\n"
        "  - 10 tools total, JSON-schema declared, dispatched via execute_tool\n"
        "  - Security fingerprinting layer: 20x-272x smaller than raw kubectl JSON\n"
        "\n"
        "Static Engine (analyzer.py)\n"
        "  - 14 deterministic checks (K8S-001..K8S-014)\n"
        "  - CHECK_REGISTRY exposes checks by ID for the agent\n"
        "\n"
        "Persistence (web/database.py)\n"
        "  - SQLite at data/kubesentinel.db\n"
        "  - Tables: users, manifests, clusters, scans, findings, images\n"
        "  - File uploads in data/uploads/manifests/, kubeconfigs in data/kubeconfigs/ (chmod 600)\n"
        "\n"
        "External Systems\n"
        "  - Anthropic API (claude-sonnet-4-6 by default)\n"
        "  - kubectl (live cluster queries, SA impersonation)\n"
        "  - helm (chart rendering)\n"
        "  - trivy (CVE scanning)\n"
        "  - GitHub API (PR comment posting from CI)"
    )

    # ── 5. Component Deep Dive ─────────────────────────────────────────────────
    add_heading(doc, "5. Component Deep Dive", level=1)

    add_heading(doc, "5.1 CLI Entry Point (agent.py)", level=2)
    add_para(doc,
        "agent.py is a thin orchestration shell. It parses CLI flags, loads .env, decides between "
        "single-path mode and multi-file (--files) mode, then dispatches to either the static "
        "engine (--no-ai) or the agentic loop. After the scan, it optionally invokes "
        "generate_patches_for_findings, applies suppressions, and renders the report (Markdown, "
        "JSON, or PR-comment format). Exit code 2 is returned when CRITICAL findings exist."
    )
    add_para(doc, "Key flags:", bold=True)
    add_bullets(doc, [
        "manifest (positional): file, directory, or Helm chart path.",
        "--files FILE [FILE...]: CI/PR mode, accepts multiple explicit files.",
        "--no-ai: skip Claude, run only the 14 static checks.",
        "--patch: post-scan, generate AI YAML patches for every finding.",
        "--model: override the Claude model (or set K8S_CHECKER_MODEL env var).",
        "--output/-o: save the rendered report to a file.",
        "--json: emit raw JSON findings to stdout.",
        "--pr-comment: format output as a compact GitHub PR comment."
    ])

    add_heading(doc, "5.2 Static Analyzer (analyzer.py)", level=2)
    add_para(doc,
        "analyzer.py owns the YAML parser and the 14-check static engine. load_manifests handles "
        "both single files (multi-doc YAML supported) and directories (recursive rglob with an "
        "exclusion list). Each check function takes (resource, context), returns either None, a "
        "single finding dict, or a list of findings. CHECK_REGISTRY maps check IDs to functions, "
        "letting the agent invoke a single check on demand via the run_check tool."
    )
    add_para(doc, "All 14 checks listed:", bold=True)
    add_bullets(doc, [
        "check_privileged_containers (K8S-001)",
        "check_host_namespace (K8S-002 - hostPID / hostIPC / hostNetwork)",
        "check_root_user (K8S-003)",
        "check_capabilities (K8S-004 - SYS_ADMIN, NET_ADMIN, SYS_PTRACE, SYS_MODULE, "
        "DAC_OVERRIDE, ALL)",
        "check_read_only_root_fs (K8S-005)",
        "check_resource_limits (K8S-006)",
        "check_image_tag (K8S-007 - :latest / untagged)",
        "check_service_account (K8S-008 - automountServiceAccountToken)",
        "check_host_path_volumes (K8S-009)",
        "check_network_policy (K8S-010 - missing labels)",
        "check_secrets_in_env (K8S-011 - heuristic on common env names)",
        "check_liveness_readiness (K8S-012)",
        "check_security_context (K8S-013 - pod-level securityContext + seccomp)",
        "check_rbac_wildcard (K8S-014)"
    ])

    add_heading(doc, "5.3 Agentic Layer (claude_agent.py)", level=2)
    add_para(doc,
        "claude_agent.py is the brain of the system. It instantiates the Anthropic client, "
        "loads the senior-Kubernetes-security-architect system prompt, and runs up to 25 "
        "iterations of the tool_use loop. On each iteration, it sends the current message history "
        "(plus tool results from the previous turn) to Claude, parses the assistant content for "
        "tool_use blocks, executes each requested tool via tools.execute_tool, and appends the "
        "tool_result back into the conversation. The loop terminates when Claude returns "
        "end_turn, when stop_reason is not tool_use, or when the finish tool flips state['done']."
    )
    add_para(doc, "Three public functions:", bold=True)
    add_bullets(doc, [
        "analyze_with_agent(manifest_path, api_key, verbose, model): single-file / directory / "
        "Helm chart scan. Picks an input hint string to guide Claude (file vs. directory vs. "
        "Helm chart).",
        "analyze_cluster_with_agent(cluster_name, kubeconfig_path, api_key, verbose, model): "
        "live-cluster scan using a stored kubeconfig. Stores the kubeconfig path in state so "
        "downstream kubectl calls can pick it up via KUBECONFIG.",
        "generate_patches_for_findings(findings, api_key, verbose, model): the post-scan patch "
        "loop. Filters out findings that already have a patch, builds a compact findings "
        "summary, and constrains the available tool list to suggest_patch and finish."
    ])
    add_para(doc,
        "Both system prompts call out explicit prompt-injection risk: a malicious YAML could "
        "embed instructions trying to override the system prompt. The tool_use API mitigates "
        "but does not eliminate this risk."
    )
    add_para(doc, "Tool call truncation:", bold=True)
    add_para(doc,
        "Tool results exceeding 8000 characters are truncated before being sent back to Claude. "
        "This protects context window utilisation against very large kubectl outputs."
    )

    add_heading(doc, "5.4 Tool Layer (tools.py)", level=2)
    add_para(doc,
        "tools.py declares the JSON schema for every tool, builds two variants of the tool "
        "list (build_tools(patch_enabled=True/False)), and contains the per-tool execution "
        "functions. The execute_tool dispatcher routes the agent's calls to the correct "
        "implementation."
    )
    add_para(doc, "Tools available to the agent:", bold=True)
    rows = [
        ("load_manifest", "Parse YAML; multi-doc and directory inputs supported."),
        ("render_helm_chart", "Run 'helm template' on a chart directory and load the rendered resources."),
        ("query_cluster", "kubectl get against an allow-listed resource type; returns compact security fingerprints."),
        ("run_check", "Execute one of the 14 static checks (or ALL) against the loaded resources."),
        ("lookup_image_cves", "Trivy scan of a single container image; cached per scan."),
        ("report_finding", "Record an AI-discovered finding with severity, attack_scenario, telco_relevance."),
        ("probe_service_account", "kubectl auth can-i --as impersonation across 13 sensitive verbs/resources."),
        ("scan_cluster_images", "Build image->pods map; Trivy-scan up to 8 unique images; correlate with risk signals."),
        ("suggest_patch", "Attach a YAML patch + explanation to a finding. Patch-loop-only tool."),
        ("finish", "Mark analysis complete; carries a summary string back to the caller."),
    ]
    add_table(doc, ["Tool", "Purpose"], rows, col_widths=[1.6, 4.7])

    add_para(doc, "Security fingerprinting layer:", bold=True)
    add_para(doc,
        "_fingerprint_cluster_resource consumes raw kubectl JSON and emits the security-relevant "
        "essence only. Containers collapse to a fixed 8-field structure (name, image, privileged, "
        "runAsUser, runAsNonRoot, readOnlyFS, allowPrivEsc, caps_add/drop). Pods get a "
        "pre-computed signals[] list (hostPID, hostIPC, hostNetwork, sa_token_automounted, "
        "privileged:<name>, root_user:<name>, writable_fs:<name>, priv_esc_allowed:<name>, "
        "caps_add[...]:<name>, hostPath:<paths>). RBAC roles get pre-parsed sensitive[] lists "
        "and wildcard_verb / wildcard_res flags. Typical size reductions: Secrets 272x, Pods 20x, "
        "RBAC roles 2x. Target: under 10 KB per query_cluster response, under $0.10 per full "
        "cluster scan."
    )
    add_para(doc, "Service Account probe matrix:", bold=True)
    add_para(doc,
        "_SA_SENSITIVE_CHECKS runs 13 'kubectl auth can-i' impersonation tests per SA: get/list "
        "secrets, get/list configmaps, create/delete pods, create/delete deployments, list "
        "serviceaccounts, create/delete clusterrolebindings, list namespaces, get nodes. System "
        "namespaces (kube-system, kube-public, kube-node-lease) are skipped to avoid noise."
    )

    add_heading(doc, "5.5 Reporter (reporter.py)", level=2)
    add_para(doc,
        "reporter.py renders findings into Markdown. Two functions: render_report (full "
        "Markdown for CLI / file output, including a Summary table, weighted Risk score, "
        "Findings detail, Remediation priority, and Suppressed findings section), and "
        "render_pr_comment (compact GitHub PR comment with collapsible 'View all findings' "
        "details block, a status banner of BLOCKED / WARNING / ADVISORY / PASSED, and an "
        "optional workflow run link). The Risk score formula is "
        "CRITICAL*10 + HIGH*5 + MEDIUM*2 + LOW*1."
    )

    add_heading(doc, "5.6 Suppressor (suppressor.py)", level=2)
    add_para(doc,
        "Looks for .k8s-checker-ignore.yaml next to the scanned manifest (and falls back to the "
        "current working directory). Each suppression rule names a check_id and optionally a "
        "resource ('Kind/name'). Suppressed findings are not removed from the report; they are "
        "split off into a separate 'Suppressed findings' section so reviewers can see the audit "
        "trail of accepted risks."
    )

    add_heading(doc, "5.7 Web Server (server.py and web/)", level=2)
    add_para(doc,
        "server.py instantiates a FastAPI app with session middleware (cookie name ks_session, "
        "key persisted in data/.secret_key chmod 600). Nine routers are wired in: setup, auth, "
        "dashboard, manifests, clusters, images, scans, users, api. A setup_guard middleware "
        "redirects every request to /setup until at least one user exists. On startup, init_db "
        "creates tables (with safe ALTER TABLE migrations for added columns), the APScheduler "
        "BackgroundScheduler starts, and any cluster-scan schedules are restored. On shutdown "
        "the scheduler is told to stop without waiting on jobs."
    )
    add_para(doc, "Notable endpoints / pages:", bold=True)
    add_bullets(doc, [
        "/setup - first-run wizard; creates the admin user.",
        "/dashboard - posture overview with severity counters and recent scans.",
        "/manifests - upload YAML, choose AI Agent or Static mode, kick off a scan.",
        "/clusters - upload a kubeconfig, optionally set a recurring schedule (6h, 12h, 24h, "
        "48h, weekly).",
        "/images - container-image CVE dashboard rolled up across scans.",
        "/scans/<id> - per-scan detail page with the 'Generate AI Patches' button.",
        "/users - admin-only user management."
    ])

    add_heading(doc, "5.8 Database Schema (web/database.py)", level=2)
    rows = [
        ("users", "id, username, email, hashed_password (bcrypt), is_admin, is_active, created_at."),
        ("manifests", "id, filename, file_path, uploaded_by/uploaded_by_name, uploaded_at."),
        ("clusters", "id, name, kubeconfig_path, added_by, schedule_hours, last_scanned_at."),
        ("scans", "id, scan_type (manifest|cluster), target_id, scan_mode (ai|static), status (queued|running|done|failed), severity counters, triggered_by, error_message, patches_status."),
        ("findings", "id, scan_id, check_id, severity, source (static|claude-ai|sa-probe|compound), context, title, detail, remediation, resource_path, attack_scenario, telco_relevance, suggested_patch, patch_explanation."),
        ("images", "id, scan_id, image_ref, critical_cves, high_cves, medium_cves, low_cves, total_cves, cve_details (JSON), scanned_at."),
    ]
    add_table(doc, ["Table", "Columns / purpose"], rows, col_widths=[1.4, 4.9])
    add_para(doc,
        "init_db() also runs _migrate() which performs idempotent ALTER TABLE ADD COLUMN "
        "statements for fields that were added after the initial schema (suggested_patch, "
        "patch_explanation, patches_status), so existing databases upgrade silently."
    )

    add_heading(doc, "5.9 Background Scanner (web/scanner.py)", level=2)
    add_para(doc,
        "run_scan(scan_id) is the FastAPI BackgroundTask that performs an asynchronous scan. It "
        "loads the Scan row, sets status='running', and dispatches based on scan_type. "
        "Manifest scans either call analyze_with_agent (AI mode) or load_manifests + "
        "run_static_checks (static mode). Cluster scans call analyze_cluster_with_agent or, "
        "in static mode, _cluster_static_checks which runs hand-rolled kubectl-based checks "
        "directly. After findings are collected, _persist_findings inserts them with per-severity "
        "counters, _extract_images deduplicates container images and Trivy-scans each one, and "
        "the scan status flips to 'done'. Exceptions are caught and surfaced through scan.error_message."
    )
    add_para(doc,
        "run_patch_generation(scan_id) is a separate post-scan flow triggered by the 'Generate "
        "AI Patches' button. It pulls existing findings out of the DB, calls "
        "generate_patches_for_findings, and writes the patch_yaml + explanation back onto each "
        "Finding row. The Scan.patches_status column tracks the lifecycle "
        "(none | generating | done | failed)."
    )
    add_para(doc, "Compound risk correlation (_correlate_risks):", bold=True)
    add_para(doc,
        "After all per-pod findings exist, this function iterates over every running pod and "
        "tallies four signal types: (1) CVE - the pod's image has CRITICAL or 5+ HIGH CVEs from "
        "Trivy, (2) Misconfiguration - the pod context has CRITICAL/HIGH static or AI findings, "
        "(3) RBAC - the pod's SA has runtime-confirmed dangerous API access from sa-probe, (4) "
        "Network - the pod's namespace has no NetworkPolicy. Pods with 2+ signals produce a "
        "CMP-001..CMP-004 compound finding with a full exploit-chain attack_scenario string."
    )

    add_heading(doc, "5.10 Scheduler (web/scheduler.py)", level=2)
    add_para(doc,
        "APScheduler BackgroundScheduler with a SQLAlchemyJobStore backed by the same engine "
        "as the application. upsert_cluster_schedule(cluster_id, name, schedule_hours, scan_mode) "
        "replaces any existing job with id 'cluster_<id>'. _scheduled_cluster_scan creates a new "
        "Scan row with triggered_by='scheduler' and invokes run_scan. restore_schedules() runs "
        "on startup to recover any schedules whose rows exist in the clusters table."
    )

    # ── 6. Two-phase design ────────────────────────────────────────────────────
    add_heading(doc, "6. The Two-Phase Scan -> Patch Design", level=1)
    add_para(doc,
        "KubeSentinel deliberately separates the scan loop from the patch loop. This is a "
        "cost-and-correctness optimisation."
    )
    add_para(doc, "Phase 1 - Scan (always runs):", bold=True)
    add_code(doc,
        "load_manifest / render_helm_chart\n"
        "    |\n"
        "    v\n"
        "run_check(ALL)             <- 14 static checks across all resources\n"
        "    |\n"
        "    v\n"
        "lookup_image_cves          <- Trivy scan per unique image\n"
        "    |\n"
        "    v\n"
        "query_cluster              <- compact security fingerprints (no raw JSON)\n"
        "    |\n"
        "    v\n"
        "probe_service_account      <- runtime SA permission proof\n"
        "    |\n"
        "    v\n"
        "scan_cluster_images        <- CVE scan on running images\n"
        "    |\n"
        "    v\n"
        "report_finding             <- AI findings + compound risk\n"
        "    |\n"
        "    v\n"
        "finish"
    )
    add_para(doc,
        "The scan loop's tool list deliberately excludes suggest_patch (build_tools(patch_enabled=False)). "
        "This keeps scan cost predictable and prevents Claude from spending tokens on patch "
        "generation for findings the user may never actually want to fix."
    )
    add_para(doc, "Phase 2 - Patch (on demand only):", bold=True)
    add_code(doc,
        "findings (from any scan - static or AI)\n"
        "    |\n"
        "    v\n"
        "suggest_patch x N         <- minimal YAML + one-sentence explanation\n"
        "    |\n"
        "    v\n"
        "finish                    <- patches written into DB / returned to CLI"
    )
    add_para(doc,
        "The patch loop uses a tightly scoped system prompt (_PATCH_GEN_SYSTEM) and a tool list "
        "filtered down to suggest_patch + finish. This prevents the model from re-running scans, "
        "re-querying kubectl, or otherwise wandering off-task while patches are being generated."
    )

    # ── 7. Compound risk ───────────────────────────────────────────────────────
    add_heading(doc, "7. Compound Risk Correlation Engine", level=1)
    add_para(doc,
        "A privileged container is bad. A wildcard ServiceAccount is bad. A namespace without "
        "a NetworkPolicy is bad. But together they are a complete exploit chain. The compound "
        "engine is what turns a list of unrelated issues into a triaged attack narrative."
    )
    rows = [
        ("CVE signal", "Image has CRITICAL CVEs or 5+ HIGH CVEs (Trivy)."),
        ("Misconfiguration signal", "Pod context has CRITICAL/HIGH static or AI findings."),
        ("RBAC signal", "SA has runtime-confirmed dangerous access (kubectl auth can-i)."),
        ("Network signal", "Pod's namespace has no NetworkPolicy coverage."),
    ]
    add_table(doc, ["Signal", "Trigger"], rows, col_widths=[1.8, 4.5])
    add_para(doc, "Severity rules:", bold=True)
    add_bullets(doc, [
        "2 signals: CMP-001",
        "3 signals: CMP-002 / CMP-003",
        "4 signals: CMP-004",
        "Severity = CRITICAL if (CVE AND RBAC) or (CVE AND misconfig AND >=3 signals); else HIGH.",
        "attack_scenario is composed as 'Attacker chain: <step 1> -> <step 2> -> ...' joining "
        "the individual signal narratives."
    ])

    # ── 8. Token efficiency ────────────────────────────────────────────────────
    add_heading(doc, "8. Token-Efficient Cluster Fingerprinting", level=1)
    add_para(doc,
        "Raw kubectl JSON is enormous: a single PodList on a non-trivial cluster easily breaks "
        "300 KB; SecretLists carry base64-encoded payloads in the hundreds of KB. Feeding that "
        "into a model context is both expensive and unhelpful (most fields are irrelevant to "
        "security). _fingerprint_cluster_resource intercepts every kubectl JSON response and "
        "emits a security-focused projection."
    )
    rows = [
        ("PodList", "~300 KB", "~6 KB", "20x to 50x"),
        ("ClusterRoleList", "~200 KB", "~5 KB", "40x"),
        ("SecretList (data stripped)", "~500 KB", "~1 KB", "272x"),
        ("RBAC binding lists", "~80 KB", "~3 KB", "25x"),
    ]
    add_table(doc, ["Resource", "Raw size", "Fingerprint", "Reduction"], rows,
              col_widths=[2.0, 1.4, 1.4, 1.5])
    add_para(doc,
        "Pre-computed signals embedded in each fingerprint (e.g. 'privileged:api', "
        "'root_user:init', 'hostPID', 'sa_token_automounted') let Claude triage without "
        "re-parsing the full container struct. Target cost: under $0.10 per full cluster scan "
        "on the default model."
    )

    # ── 9. Check catalogue ─────────────────────────────────────────────────────
    add_heading(doc, "9. Static Check Catalogue (K8S-001 through K8S-014)", level=1)
    rows = [
        ("K8S-001", "Privileged container", "CRITICAL", "securityContext.privileged: true"),
        ("K8S-002", "Host namespaces", "CRITICAL / HIGH", "spec.hostPID / hostIPC / hostNetwork"),
        ("K8S-003", "Root user", "HIGH / MEDIUM", "runAsUser: 0 or runAsNonRoot: false"),
        ("K8S-004", "Dangerous capabilities", "CRITICAL / HIGH", "capabilities.add includes SYS_ADMIN / ALL etc."),
        ("K8S-005", "Writable root filesystem", "MEDIUM", "readOnlyRootFilesystem not true"),
        ("K8S-006", "Missing resource limits / requests", "MEDIUM / LOW", "resources.limits or resources.requests unset"),
        ("K8S-007", "Unpinned image tag", "MEDIUM", "image uses :latest or no tag"),
        ("K8S-008", "SA token automount", "MEDIUM", "automountServiceAccountToken not false"),
        ("K8S-009", "hostPath volumes", "CRITICAL / HIGH", "/, /etc, /proc, /sys, docker.sock -> CRITICAL"),
        ("K8S-010", "Missing labels", "LOW", "metadata.labels empty (impairs NetworkPolicy targeting)"),
        ("K8S-011", "Hardcoded secrets in env", "HIGH", "env name contains password/secret/token/key + literal value"),
        ("K8S-012", "Missing probes", "LOW", "livenessProbe or readinessProbe absent"),
        ("K8S-013", "Pod-level securityContext / seccomp", "MEDIUM / LOW", "no pod securityContext or no seccompProfile"),
        ("K8S-014", "RBAC wildcard verbs / resources", "CRITICAL / HIGH", "verbs: ['*'] or resources: ['*']"),
    ]
    add_table(doc, ["Check ID", "Title", "Severity", "Trigger condition"], rows,
              col_widths=[0.9, 2.1, 1.2, 2.5])
    add_para(doc, "Beyond the 14:", bold=True)
    add_bullets(doc, [
        "SAP-001 - runtime ServiceAccount probe finding, raised by sa-probe.",
        "CMP-001..CMP-004 - compound risk findings raised by the correlation engine.",
        "AI-* - dynamically chosen by Claude through report_finding for issues that fall "
        "outside the deterministic check set (logic-level misconfigurations, supply chain "
        "concerns, telco/CNF specifics, etc.)."
    ])

    # ── 10. Demo ───────────────────────────────────────────────────────────────
    add_heading(doc, "10. Demonstration Guide", level=1)
    add_para(doc,
        "This walkthrough validates every layer of KubeSentinel - from unit tests through "
        "AI scanning to the web dashboard. Run each step from the project root with the venv "
        "activated."
    )

    add_heading(doc, "Step 0 - One-time setup", level=2)
    add_code(doc,
        "git clone https://github.com/jaydenaung/kubesentinel.git\n"
        "cd kubesentinel\n"
        "python3 -m venv venv\n"
        "source venv/bin/activate                  # macOS/Linux\n"
        "# venv\\Scripts\\activate                  # Windows\n"
        "python -m pip install -r requirements.txt\n"
        "cp .env.example .env\n"
        "# edit .env and set ANTHROPIC_API_KEY=sk-ant-..."
    )
    add_para(doc, "Optional binaries (KubeSentinel skips gracefully if absent):", bold=True)
    add_code(doc,
        "brew install trivy      # macOS - CVE scanning\n"
        "brew install helm       # Helm chart rendering\n"
        "# kubectl via your cloud provider CLI or kubernetes.io docs"
    )

    add_heading(doc, "Step 1 - Run the unit tests", level=2)
    add_code(doc, "pytest tests/ -v")
    add_para(doc,
        "Expected: 48 tests pass (40 in test_analyzer.py covering all 14 static checks, "
        "8 in test_suppressor.py). No API key or cluster required."
    )

    add_heading(doc, "Step 2 - Static scan (no API key)", level=2)
    add_code(doc, "python agent.py samples/vulnerable.yaml --no-ai")
    add_para(doc,
        "Expect 10+ findings across CRITICAL and HIGH severities on the intentionally "
        "misconfigured sample manifest. Exit code will be 2 due to CRITICAL findings - this is "
        "what CI uses to block merges."
    )

    add_heading(doc, "Step 3 - AI agent scan", level=2)
    add_code(doc,
        "export ANTHROPIC_API_KEY=sk-ant-your-key-here\n"
        "python agent.py samples/vulnerable.yaml"
    )
    add_para(doc, "You will see tool calls scroll past, e.g.:", bold=True)
    add_code(doc,
        "      load_manifest(path=samples/vulnerable.yaml)\n"
        "      run_check(check=ALL  resource=-1)\n"
        "      lookup_image_cves(nginx:latest)\n"
        "      query_cluster(pods)\n"
        "      probe_service_account(vulnerable-sa  default)\n"
        "      report_finding([CRITICAL] Privileged container with host namespaces ...)\n"
        "      finish(Found 14 findings ...)"
    )
    add_para(doc,
        "Notice suggest_patch does not appear - patch generation is a separate step."
    )

    add_heading(doc, "Step 4 - Generate AI patches (premium)", level=2)
    add_code(doc,
        "# AI scan + AI patches\n"
        "python agent.py samples/vulnerable.yaml --patch\n"
        "\n"
        "# Static scan + AI patches (works on any scan type)\n"
        "python agent.py samples/vulnerable.yaml --no-ai --patch\n"
        "\n"
        "# Inspect the JSON form\n"
        "python agent.py samples/vulnerable.yaml --patch --json | python -m json.tool"
    )

    add_heading(doc, "Step 5 - Web dashboard end-to-end", level=2)
    add_code(doc, "python server.py")
    add_bullets(doc, [
        "Open http://localhost:8000 and complete the setup wizard.",
        "Manifests -> Upload Manifest -> select samples/vulnerable.yaml.",
        "Choose AI Agent or Static, click Upload & Scan.",
        "Watch the status badge cycle: queued -> running -> done.",
        "Click into the scan to see findings, severity breakdown, and compound risk sections.",
        "Click 'Generate AI Patches' (top right) - patches appear inline per finding.",
        "Navigate to Images to inspect CVE counts (Trivy must be installed)."
    ])

    add_heading(doc, "Step 6 - Static-only dashboard scan", level=2)
    add_para(doc,
        "Repeat Step 5 with 'Static' mode selected to demonstrate the air-gapped path. "
        "No API key needed; findings appear in seconds."
    )

    add_heading(doc, "Step 7 - PR-level scanning via GitHub Actions", level=2)
    add_para(doc,
        "Copy .github/workflows/kubesentinel.yml into your target repo, add ANTHROPIC_API_KEY "
        "as a repository secret, and open a PR that modifies any .yaml file. The workflow "
        "scans only the changed files, posts findings as a PR comment, and fails the check on "
        "CRITICAL findings. To dry-run locally:"
    )
    add_code(doc,
        "python agent.py --files samples/vulnerable.yaml --pr-comment"
    )

    add_heading(doc, "Step 8 - Suppression allowlist", level=2)
    add_code(doc,
        "cp samples/.k8s-checker-ignore.yaml .\n"
        "python agent.py samples/vulnerable.yaml --no-ai"
    )
    add_para(doc,
        "Suppressed findings still appear in the report footer with their stated reason, "
        "providing an audit trail for compliance reviews."
    )

    # ── 11. Start/Stop ─────────────────────────────────────────────────────────
    add_heading(doc, "11. Starting and Stopping the Server", level=1)
    add_para(doc, "Foreground start (recommended for demos):", bold=True)
    add_code(doc,
        "source venv/bin/activate\n"
        "python server.py                    # binds 0.0.0.0:8000\n"
        "python server.py --port 8080        # custom port\n"
        "python server.py --host 127.0.0.1   # local-only\n"
        "# Press Ctrl+C to stop"
    )
    add_para(doc, "Background start (long-running deployment):", bold=True)
    add_code(doc,
        "nohup python server.py > kubesentinel.log 2>&1 &\n"
        "echo $! > kubesentinel.pid\n"
        "\n"
        "# Tail the log\n"
        "tail -f kubesentinel.log"
    )
    add_para(doc, "Stop a background server:", bold=True)
    add_code(doc,
        "# By PID file\n"
        "kill $(cat kubesentinel.pid) && rm kubesentinel.pid\n"
        "\n"
        "# Or by port\n"
        "lsof -ti:8000 | xargs kill -9"
    )
    add_para(doc, "Production hardening checklist:", bold=True)
    add_bullets(doc, [
        "Reverse-proxy behind nginx or Caddy with TLS termination.",
        "Set ANTHROPIC_API_KEY via the system service manager (systemd EnvironmentFile, "
        "Docker secret, etc.), not the .env file.",
        "Restrict --host to an internal IP and rely on the proxy for external access.",
        "Run the process under a dedicated user with read/write on data/ and nothing else.",
        "Back up data/kubesentinel.db on whatever cadence matches your retention policy.",
        "Keep kubeconfigs at chmod 600 (the server already enforces this on upload via the "
        "data/kubeconfigs/ directory)."
    ])
    add_para(doc, "Graceful shutdown internals:", bold=True)
    add_para(doc,
        "FastAPI's @app.on_event('shutdown') handler calls scheduler.shutdown(wait=False), "
        "which cancels any pending scheduled scan and lets running BackgroundTasks finish their "
        "current work. SQLite writes are flushed when the SQLAlchemy session closes. After Ctrl+C, "
        "the process exits cleanly with no orphaned files."
    )

    # ── 12. CLI reference ─────────────────────────────────────────────────────
    add_heading(doc, "12. CLI Reference", level=1)
    rows = [
        ("python agent.py samples/vulnerable.yaml", "AI agent scan on a single file."),
        ("python agent.py samples/vulnerable.yaml --no-ai", "Static checks only (no API key needed)."),
        ("python agent.py samples/vulnerable.yaml --patch", "AI scan + post-scan patch generation."),
        ("python agent.py samples/vulnerable.yaml --no-ai --patch", "Static scan + AI patches."),
        ("python agent.py k8s/", "Scan all YAML files in a directory recursively."),
        ("python agent.py ./my-helm-chart/", "Detect Chart.yaml and render with helm template first."),
        ("python agent.py file.yaml --output reports/result.md", "Persist a Markdown report."),
        ("python agent.py file.yaml --patch --json", "Raw JSON output (good for piping)."),
        ("python agent.py --files a.yaml b.yaml --pr-comment", "CI/PR mode with compact PR comment formatting."),
        ("python agent.py file.yaml --model claude-haiku-4-5-20251001", "Override the model used."),
    ]
    add_table(doc, ["Command", "Effect"], rows, col_widths=[3.6, 2.7])
    add_para(doc, "Exit codes:", bold=True)
    add_bullets(doc, [
        "0 = clean (no findings or only LOW/MEDIUM/INFO)",
        "1 = error (missing path, missing API key in non-fallback mode)",
        "2 = CRITICAL findings detected (suitable as a CI gating signal)"
    ])

    # ── 13. Configuration ─────────────────────────────────────────────────────
    add_heading(doc, "13. Configuration and Environment", level=1)
    rows = [
        ("ANTHROPIC_API_KEY", ".env or shell export", "Required for AI mode and patch generation."),
        ("K8S_CHECKER_MODEL", "env var", "Override Claude model (default claude-sonnet-4-6)."),
        ("--model FLAG", "CLI flag", "Per-invocation model override; wins over env var."),
        ("GITHUB_SERVER_URL / GITHUB_REPOSITORY / GITHUB_RUN_ID", "env vars (CI)", "Used to build the PR-comment workflow link."),
        ("KUBECONFIG", "env var (auto)", "Set automatically when a stored kubeconfig is selected."),
    ]
    add_table(doc, ["Variable", "Source", "Purpose"], rows, col_widths=[2.4, 1.6, 2.4])
    add_para(doc, "Filesystem locations:", bold=True)
    add_bullets(doc, [
        "data/kubesentinel.db - SQLite database (users, manifests, clusters, scans, findings, images).",
        "data/.secret_key - 64-char hex session secret (auto-generated, chmod 600).",
        "data/uploads/manifests/ - uploaded YAML files.",
        "data/kubeconfigs/ - uploaded kubeconfigs (chmod 600 enforced).",
        ".env - per-project ANTHROPIC_API_KEY (loaded by python-dotenv).",
        ".k8s-checker-ignore.yaml - suppression allowlist (searched next to manifest then cwd)."
    ])

    # ── 14. Security / limitations ────────────────────────────────────────────
    add_heading(doc, "14. Security, Trust, and Limitations", level=1)
    add_para(doc, "Trust model:", bold=True)
    add_bullets(doc, [
        "Read-only: KubeSentinel never patches your cluster, never mutates manifests, never "
        "writes anything outside the project's data/ directory and any --output path you set.",
        "Local-first: SQLite + filesystem + bcrypt session cookies. No external SaaS dependency. "
        "The Anthropic API is the only outbound call, and only when AI mode is active.",
        "Auditable: every finding has a check_id, source field (static / claude-ai / sa-probe / "
        "compound), and (for AI findings) an attack_scenario string. Patches carry a "
        "patch_explanation."
    ])
    add_para(doc, "Known limitations:", bold=True)
    add_bullets(doc, [
        "Prompt-injection risk: malicious YAML could embed text trying to override the system "
        "prompt. tool_use mitigates this but does not eliminate it. Never scan untrusted YAML.",
        "AI findings can hallucinate: anything tagged [AI] must be independently reviewed.",
        "Patches are minimal snippets: they show the changed field(s) with parent-key context "
        "but do not produce a full corrected manifest. Apply with human judgment.",
        "Trivy / kubectl / helm are optional: features depending on them silently skip when the "
        "binary is missing.",
        "scan_cluster_images caps at 8 unique images per call to avoid runaway scanning on large "
        "clusters.",
        "_SA_SENSITIVE_CHECKS is opinionated - it covers 13 high-value verb/resource pairs, not "
        "the full RBAC surface."
    ])
    add_para(doc, "Disclaimer:", bold=True)
    add_para(doc,
        "KubeSentinel is provided for informational and educational purposes only, without "
        "warranty. A clean report does not mean your cluster is secure. Combine with manual "
        "review, penetration testing, and defence-in-depth.",
        italic=True
    )

    # ── 15. Project layout ────────────────────────────────────────────────────
    add_heading(doc, "15. Project Layout", level=1)
    add_code(doc,
        "kubesentinel/\n"
        "  agent.py              CLI entry point\n"
        "  analyzer.py           YAML parser + 14 static checks + CHECK_REGISTRY\n"
        "  claude_agent.py       Agentic loop using Anthropic tool_use API\n"
        "  tools.py              Tool schemas + execution + security fingerprinting\n"
        "  reporter.py           Markdown / PR-comment rendering\n"
        "  suppressor.py         .k8s-checker-ignore.yaml loader + filter\n"
        "  server.py             FastAPI server entry point\n"
        "  requirements.txt\n"
        "  .env.example\n"
        "  .github/workflows/\n"
        "      kubesentinel.yml  PR-level manifest scanning workflow\n"
        "  web/\n"
        "      database.py       SQLAlchemy models (User, Manifest, Cluster, Scan, Finding, Image)\n"
        "      auth.py           Session auth + bcrypt password hashing\n"
        "      scanner.py        BackgroundTask scan runners (manifest + cluster)\n"
        "      scheduler.py      APScheduler integration for recurring cluster scans\n"
        "      routes/           FastAPI routers (dashboard, manifests, clusters, images, ...)\n"
        "      templates/        Jinja2 dark-theme dashboard UI\n"
        "  tests/\n"
        "      test_analyzer.py     40 unit tests covering all 14 static checks\n"
        "      test_suppressor.py    8 unit tests for suppression logic\n"
        "  samples/\n"
        "      vulnerable.yaml             intentionally misconfigured manifest\n"
        "      secure.yaml                 hardened reference manifest\n"
        "      test-sa-probe.yaml          SA probe + compound risk test manifest\n"
        "      .k8s-checker-ignore.yaml    example suppression config\n"
        "  data/                runtime data (gitignored)"
    )

    # ── 16. Roadmap ────────────────────────────────────────────────────────────
    add_heading(doc, "16. Roadmap", level=1)
    rows = [
        ("1",  "Shipped", "AI patch generation - --patch CLI flag + 'Generate AI Patches' button."),
        ("1b", "Shipped", "Runtime SA probing + compound risk correlation (CVE+misconfig+RBAC+network)."),
        ("1c", "Shipped", "Token-efficient cluster fingerprinting (20x to 272x smaller than raw JSON)."),
        ("2",  "In progress", "Patch review UI - diff viewer with approve / reject workflow."),
        ("3",  "Planned", "GitHub PR creation - agent opens fix PRs after human approval."),
        ("4",  "Planned", "Compliance report generator - map findings to CIS / NIST / SOC2 controls."),
        ("5",  "Planned", "Findings relationship graph - model attack paths across CVE -> image -> deployment -> RBAC."),
        ("6",  "Planned", "Runtime signals - Falco / Kubernetes audit log integration."),
        ("7",  "Planned", "Multi-agent architecture - triage, remediation, compliance, orchestrator agents."),
    ]
    add_table(doc, ["Phase", "Status", "Feature"], rows, col_widths=[0.7, 1.1, 4.5])

    # ── Footer ────────────────────────────────────────────────────────────────
    doc.add_paragraph()
    foot = doc.add_paragraph()
    foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = foot.add_run(
        "KubeSentinel - AI-Powered Kubernetes Security Platform | "
        "Copyright 2026 Jayden Aung | Apache License 2.0"
    )
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    output_path = "KubeSentinel_Technical_Documentation.docx"
    doc.save(output_path)
    print(f"Wrote: {output_path}")


if __name__ == "__main__":
    main()
